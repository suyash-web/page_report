from pydoc import doc
from turtle import width
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
import Compressor
import os

cwd = os.getcwd()
mobile_images = cwd+"/output/Mobile"
desktop_images = cwd+"/output/Desktop"
def create_document_1(url :str, bot_time, scores, res):
    #cwd = os.getcwd()
    document = Document()
    head = document.add_heading()
    head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    head.add_run().add_picture("properologo.png", width=Inches(2.0), height=Inches(0.5))
    head.paragraph_format.space_after = Pt(15)
    head2 = document.add_heading()
    head2.add_run('WEBSITE  ANALYSIS  REPORT').bold = True
    head2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head2.style = document.styles['Normal']
    font_head = head2.style.font
    font_head.size = Pt(16)
    head.paragraph_format.space_before = Pt(3)
    p1 = document.add_paragraph()
    p1.add_run("Website  under  analysis:-  ").bold = True
    p1.add_run(url).italic = True
    p1.paragraph_format.space_after = Pt(60)

    p2 = document.add_paragraph()
    p2.add_run("Analyzer  1:- ").bold = True
    p2.add_run("GTmetrix").italic = True
    p2.add_run().add_break()
    p2.add_run("Results:  ").bold = True
    p2.add_run("GTmetrix  Grade  and  Web  Vitals").italic = True
    r1 = p2.add_run()
    r1.add_break()
    p2.paragraph_format.space_after = Pt(20)
    r1.add_picture(cwd+'/output/Gtmetrix/screenshot_1.png', width=Inches(6.0), height=Inches(2.0))
    r1.add_break()
    p2.paragraph_format.space_after = Pt(20)
    r1.add_picture(cwd+'/output/Gtmetrix/screenshot_2.png', width=Inches(6.0), height=Inches(2.0))
    p2.paragraph_format.space_after = Pt(60)
    document.add_page_break()

    p3 = document.add_paragraph()
    p3.add_run("Analyzer  2:-  ").bold = True
    p3.add_run("Shopify  Speedbooster").italic = True
    p3.add_run().add_break()
    p3.add_run("Results:  ").bold = True
    p3.add_run("There  are  no  oversized  images  found.").italic = True
    p3.paragraph_format.space_after = Pt(40)

    p4 = document.add_paragraph()
    p4.add_run("Analyzer  3:-  ").bold = True
    p4.add_run("PageSpeed Insights by Google Developers").italic = True
    p4.add_run().add_break()
    p4.add_run("Results:  ").bold = True
    p4.add_run().add_break()
    p4.add_run("(1.)  For  Mobile  site:  ").bold = True
    p4.add_run().add_break()
    p4.add_run("Score  :  "+ str(scores[0]) + "/100")
    p4.paragraph_format.space_after = Pt(60)
    p4.add_run().add_break()
    p4.add_run("Webvitals and Opportunities (if any):-").italic = True
    for img in os.listdir(cwd+"/output/Mobile"):
        p4.add_run().add_picture(os.path.join(mobile_images, img), width=Inches(6.0), height=Inches(2.0))
        p4.paragraph_format.space_after = Pt(20)
        p4.add_run().add_break()
    p4.add_run().add_break()
    document.add_page_break()

    p5 = document.add_paragraph()
    p5.add_run("(2.)  For  Desktop  site:  ").bold = True
    p5.add_run().add_break()
    p5.add_run("Score  :  "+ str(scores[1]) + "/100")
    p5.paragraph_format.space_after = Pt(60)
    p5.add_run().add_break()
    p5.add_run("Webvitals and Opportunities (if any):-").italic = True
    for img in os.listdir(cwd+"/output/Desktop"):
        p5.add_run().add_picture(os.path.join(desktop_images, img), width=Inches(6.0), height=Inches(2.0))
        p5.paragraph_format.space_after = Pt(20)
        p5.add_run().add_break()
    p5.add_run().add_break()
    document.add_page_break()

    p6 = document.add_paragraph()
    p6.add_run("Analyzer  4:-  ").bold = True
    p6.add_run("Solarwinds  Pingdom").italic = True
    p6.add_run().add_break()
    p6.add_run("Performance  Score:-  "+res[0])
    p6.add_run().add_break()
    p6.add_run("Page  Size:-  "+res[1])
    p6.add_run().add_break()
    p6.add_run("Load  Time:-  "+res[2])
    p6.add_run().add_break()
    p6.add_run("Requests:-  "+res[3])
    p6.paragraph_format.space_after = Pt(60)

    p7 = document.add_paragraph()
    p7.add_run("Analysis  completed  in:-  ").bold = True
    p7.add_run(str(bot_time)+" sec(s)").italic = True

    if scores[2] > "2.5 s":
        p8 = document.add_paragraph()
        p8.add_run("Your LCP score is high! Here is what you can do:- ").bold = True
        p9 = document.add_paragraph()
        p9.add_run("(1.) Reduce the impact of third-party code.")
        p10 = document.add_paragraph()
        p10.add_run("(2.) Minimize main-thread work.")
        p11 = document.add_paragraph()
        p11.add_run("(3.) Reduce the impact of third-party code.")
        p12 = document.add_paragraph()
        p12.add_run("(4.) Avoid render-blocking elements.")
        p13 = document.add_paragraph()
        p13.add_run("(5.) Avoid putting large elements like images, videos, ads, slideshows, iframes in above the fold area.")
        p14 = document.add_paragraph()
        p14.add_run("(6.) Minify CSS and JavaScript.")
        p15 = document.add_paragraph()
        p15.add_run("(7.) Avoid lazy loading of elements in above the fold area.")
    name = url[12:]
    f = name.split(".")
    document.save(cwd+'/output/'+f[0]+'.docx')

def create_document_2(url : str, bot_time, scores, res):
    files = 0
    dir = cwd+"/output/oversized_images"
    for path in os.listdir(dir):
        if os.path.isfile(os.path.join(dir, path)):
            files += 1
    document = Document()
    head = document.add_heading()
    head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    head.add_run().add_picture("properologo.png", width=Inches(2.0), height=Inches(0.5))
    head.paragraph_format.space_after = Pt(15)
    head2 = document.add_heading()
    head2.add_run('WEBSITE  ANALYSIS  REPORT').bold = True
    head2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    head2.style = document.styles['Normal']
    font_head = head2.style.font
    font_head.size = Pt(16)
    head.paragraph_format.space_before = Pt(3)
    p1 = document.add_paragraph()
    p1.add_run("Website  under  analysis:-  ").bold = True
    p1.add_run(url).italic = True
    p1.paragraph_format.space_after = Pt(60)

    p2 = document.add_paragraph()
    p2.add_run("Analyzer 1:- ").bold = True
    p2.add_run("GTmetrix").italic = True
    p2.add_run().add_break()
    p2.add_run("Results:  ").bold = True
    p2.add_run("GTmetrix  Grade  and  Web  Vitals").italic = True
    p2.add_run().add_break()
    p2.paragraph_format.space_after = Pt(20)
    p2.add_run().add_picture(cwd+'/output/Gtmetrix/screenshot_1.png', width=Inches(6.0), height=Inches(1.5))
    p2.add_run().add_break()
    p2.paragraph_format.space_after = Pt(20)
    p2.add_run().add_picture(cwd+'/output/Gtmetrix/screenshot_2.png', width=Inches(6.0), height=Inches(1.5))
    p2.paragraph_format.space_after = Pt(60)
    document.add_page_break()

    p3 = document.add_paragraph()
    p3.add_run("Analyzer 2:- ").bold = True
    p3.add_run("Shopify  Speedbooster").italic = True
    p3.add_run().add_break()
    p3.add_run("Results:  ").bold = True
    p3.add_run("Oversized Image(s)").italic = True
    p3.paragraph_format.space_after = Pt(20)
    for i in os.listdir(cwd+"/output/oversized_images/"):
        p3.add_run().add_picture(cwd+"/output/oversized_images/"+i, width=Inches(6.0), height=Inches(2.0))
        p3.add_run().add_break()
        p3.paragraph_format.space_after = Pt(20)
    p3.paragraph_format.space_after = Pt(60)

    p5 = document.add_paragraph()
    bits = Compressor.compress_images()
    p5.add_run("Bot's  Compressor: ").bold = True
    p5.add_run("Compressed  Image(s)").italic = True
    li = os.listdir(cwd+"/output/Compressed_Images")
    for img in li:
        r4 = p5.add_run()
        r4.add_break()
        if img.endswith("jpeg"):
            r4.add_picture(cwd+'/output/Compressed_Images/'+img, width=Inches(6.0), height=Inches(2.0))
            p5.paragraph_format.space_after = Pt(20)
    p5.paragraph_format.space_after = Pt(100)
    p5.add_run().add_break()
    p5.paragraph_format.space_after = Pt(30)

    p6 = document.add_paragraph()
    p6.add_run("Result:  ").bold = True
    p6.add_run("Total  bits  saved  after  Lossless  Compression of the  images :-  "+str(bits)+"KB.").italic = True
    r5 = p5.add_run()
    r5.add_break()
    p6.paragraph_format.space_after = Pt(20)
    document.add_page_break()

    p7 = document.add_paragraph()
    p7.add_run("Analyzer  3:-  ").bold = True
    p7.add_run("PageSpeed Insights by Google Developers").italic = True
    p7.add_run().add_break()
    p7.add_run("Results:  ").bold = True
    p7.add_run().add_break()
    p7.add_run("(1.)  For  Mobile  site:  ").bold = True
    p7.add_run().add_break()
    p7.add_run("Score  :  "+ str(scores[0]) + "/100")
    p7.paragraph_format.space_after = Pt(60)
    p7.add_run().add_break()
    p7.add_run("Webvitals and Opportunities (if any):-").italic = True
    for img in os.listdir(cwd+"/output/Mobile"):
        p7.add_run().add_picture(os.path.join(mobile_images, img), width=Inches(6.0), height=Inches(2.0))
        p7.paragraph_format.space_after = Pt(20)
        p7.add_run().add_break()
    document.add_page_break()

    p8 = document.add_paragraph()
    p8.add_run("(2.)  For  Desktop  site:  ").bold = True
    p8.add_run().add_break()
    p8.add_run("Score  :  "+ str(scores[1]) + "/100")
    p8.paragraph_format.space_after = Pt(60)
    p8.add_run().add_break()
    p8.add_run("Webvitals and Opportunities (if any):-").italic = True
    for img in os.listdir(cwd+"/output/Desktop"):
        p8.add_run().add_picture(os.path.join(desktop_images, img), width=Inches(6.0), height=Inches(2.0))
        p7.paragraph_format.space_after = Pt(20)
        p8.add_run().add_break()
    document.add_page_break()

    p9 = document.add_paragraph()
    p9.add_run("Analyzer  4:-  ").bold = True
    p9.add_run("Solarwinds  Pingdom").italic = True
    p9.add_run().add_break()
    p9.add_run("Performance  Score:-  "+res[0])
    p9.add_run().add_break()
    p9.add_run("Page  Size:-  "+res[1])
    p9.add_run().add_break()
    p9.add_run("Load  Time:-  "+res[2])
    p9.add_run().add_break()
    p9.add_run("Requests:-  "+res[3])
    p9.paragraph_format.space_after = Pt(60)

    large_images = len(os.listdir(cwd+"/output/oversized_images"))
    if scores[2] > "2.5 s":
        p10 = document.add_paragraph()
        p10.add_run("Your LCP score is high! It makes your website run slow. Here is what you can do:- ").bold = True
        p11 = document.add_paragraph()
        p11.add_run("(1.) You have "+str(large_images)+" oversized images. Use the compressed ones provided.")
        p12 = document.add_paragraph()
        p12.add_run("(2.) Reduce the impact of third-party code.")
        p13 = document.add_paragraph()
        p13.add_run("(3.) Minimize main-thread work.")
        p14 = document.add_paragraph()
        p14.add_run("(4.) Reduce server response time.")
        p15 = document.add_paragraph()
        p15.add_run("(5.) Avoid render-blocking elements.")
        p16 = document.add_paragraph()
        p16.add_run("(6.) Avoid putting large elements like images, videos, ads, slideshows, iframes in above the fold area.")
        p17 = document.add_paragraph()
        p17.add_run("(7.) Minify CSS and JavaScript.")
        p18 = document.add_paragraph()
        p18.add_run("(8.) Avoid lazy loading of elements in above the fold area.")
    
    p19 = document.add_paragraph()
    p19.add_run("Analysis  completed  in:-  ").bold = True
    p19.add_run(str(bot_time)+" sec(s)").italic = True

    name = url[12:]
    f = name.split(".")
    document.save(cwd+'/output/'+f[0]+'.docx')


if __name__ == "__main__":
    url = "https://www.arkoflooring.com/"
    create_document_2(url, 200, ['21', '62', '3 s'], ['66', '3.3', '2.34', '141'])